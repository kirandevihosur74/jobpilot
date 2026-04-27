import logging
import os
import json
import re
import shutil
from pathlib import Path
from fastapi import APIRouter, HTTPException, UploadFile, File, Form, Depends
from fastapi.responses import FileResponse
import openai
from docx import Document
import io
from sqlalchemy.orm import Session
from database import get_db
from models import UserPrefs, Resume

logger = logging.getLogger("jobpilot.resume")
router = APIRouter(prefix="/api/resume", tags=["resume"])

# Resume upload directory
UPLOAD_DIR = Path(__file__).parent.parent / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)
ALLOWED_EXT = {".pdf", ".docx", ".doc", ".txt"}

_client = None

def get_client():
    global _client
    if _client is None:
        key = os.getenv("TOKENROUTER_API_KEY")
        base_url = os.getenv("TOKENROUTER_BASE_URL", "https://api.tokenrouter.com/v1")
        if not key:
            raise HTTPException(status_code=500, detail="TOKENROUTER_API_KEY not configured")
        _client = openai.OpenAI(api_key=key, base_url=base_url)
    return _client

def extract_text(file: UploadFile) -> str:
    content = file.file.read()
    if file.filename.endswith(".docx"):
        doc = Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return content.decode("utf-8", errors="ignore")

@router.post("/tailor")
async def tailor_resume(
    file: UploadFile = File(...),
    job: str = Form(...),
    prefs: str = Form(...),
):
    try:
        job_data = json.loads(job)
        prefs_data = json.loads(prefs)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid job or prefs JSON")

    resume_text = extract_text(file)
    if not resume_text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from file")

    prompt = f"""You are an expert resume coach. Tailor this resume for the specific job below.

JOB:
Title: {job_data.get('title', '')}
Company: {job_data.get('company', '')}
Location: {job_data.get('location', '')}
Description: {job_data.get('description', '')[:3000]}

CANDIDATE RESUME:
{resume_text[:4000]}

Rewrite the resume tailored to this job. Rules:
- Keep all real facts, companies, dates, metrics — never invent anything
- Reorder skills by relevance to this job
- Rewrite bullet points to emphasize relevant experience using job keywords
- Rewrite summary to target this specific role and company
- Keep same sections as original resume

Return ONLY raw JSON (no markdown):
{{
  "summary": "rewritten professional summary (3-4 sentences)",
  "skills": ["skill1", "skill2", ...],
  "experience_bullets": {{
    "most recent company/role": ["bullet1", "bullet2", "bullet3"],
    "second company/role": ["bullet1", "bullet2"]
  }},
  "keywords_added": ["keyword1", "keyword2"],
  "match_score": 8,
  "tailoring_notes": "2-3 sentences on what was changed and why"
}}"""

    try:
        message = get_client().chat.completions.create(
            model="anthropic/claude-sonnet-4.6",
            max_tokens=2048,
            messages=[{"role": "user", "content": prompt}],
        )
        text = message.choices[0].message.content or ""
        logger.info("LLM raw response (first 500): %s", text[:500])
        # strip markdown fences robustly
        import re
        text = re.sub(r"```(?:json)?\s*", "", text).replace("```", "").strip()
        # extract first {...} block
        m = re.search(r"\{[\s\S]*\}", text)
        if not m:
            logger.error("No JSON object found in: %s", text[:300])
            raise HTTPException(status_code=422, detail="No JSON object in LLM response")
        try:
            result = json.loads(m.group())
        except json.JSONDecodeError as e:
            logger.error("JSON parse failed: %s\nText: %s", e, m.group()[:300])
            raise HTTPException(status_code=422, detail=f"JSON parse failed: {e}")
        return {"tailored": result, "original_length": len(resume_text)}
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Resume tailor unexpected error: %s", e)
        raise HTTPException(status_code=500, detail=str(e))
    except openai.APIError as e:
        logger.error("TokenRouter API error (resume tailor): %s", e)
        raise HTTPException(status_code=502, detail=str(e))


def _extract_pdf_text(path: Path) -> str:
    """Extract text from PDF for resume_context auto-population."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        return "\n".join(p.extract_text() or "" for p in reader.pages)
    except Exception as e:
        logger.warning("PDF text extract failed: %s", e)
        return ""


@router.post("/upload")
async def upload_resume(file: UploadFile = File(...), db: Session = Depends(get_db)):
    """Upload + persist resume file. Stores path in UserPrefs, returns extracted text."""
    ext = Path(file.filename or "").suffix.lower()
    if ext not in ALLOWED_EXT:
        raise HTTPException(status_code=400, detail=f"Unsupported file type. Allowed: {ALLOWED_EXT}")

    dest_path = UPLOAD_DIR / f"resume{ext}"
    # Remove any old resume regardless of extension
    for existing in UPLOAD_DIR.glob("resume.*"):
        try: existing.unlink()
        except Exception: pass

    with dest_path.open("wb") as f:
        shutil.copyfileobj(file.file, f)
    logger.info("Resume saved: %s (%d bytes)", dest_path, dest_path.stat().st_size)

    # Extract text for context
    text = ""
    if ext == ".pdf":
        text = _extract_pdf_text(dest_path)
    elif ext == ".docx":
        try:
            doc = Document(str(dest_path))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            logger.warning("docx extract failed: %s", e)
    elif ext == ".txt":
        text = dest_path.read_text(errors="ignore")

    # Save metadata to UserPrefs
    row = db.query(UserPrefs).first()
    if not row:
        row = UserPrefs()
        db.add(row)
    row.resume_filename = file.filename or f"resume{ext}"
    row.resume_file_path = str(dest_path.absolute())
    if text and not row.resume_context:
        row.resume_context = text[:8000]
    db.commit()

    return {
        "filename": row.resume_filename,
        "path": row.resume_file_path,
        "size": dest_path.stat().st_size,
        "extracted_text_length": len(text),
        "extracted_text": text[:500],
    }


@router.get("/info")
def get_resume_info(db: Session = Depends(get_db)):
    """Return current resume filename + path (no file content)."""
    row = db.query(UserPrefs).first()
    if not row or not row.resume_file_path:
        return {"uploaded": False, "filename": None, "path": None}
    exists = Path(row.resume_file_path).exists()
    return {
        "uploaded": exists,
        "filename": row.resume_filename,
        "path": row.resume_file_path if exists else None,
    }


@router.delete("/upload")
def delete_resume(db: Session = Depends(get_db)):
    row = db.query(UserPrefs).first()
    if row and row.resume_file_path:
        try: Path(row.resume_file_path).unlink()
        except Exception: pass
        row.resume_filename = ""
        row.resume_file_path = ""
        db.commit()
    return {"ok": True}


# ── DOCX placeholder tailoring (Option A — preserves Google Docs formatting) ─
TAILORED_DIR = UPLOAD_DIR / "tailored"
TAILORED_DIR.mkdir(exist_ok=True)


def _slugify(s: str) -> str:
    import re
    return re.sub(r"[^a-z0-9]+", "_", (s or "").lower()).strip("_")[:60]


def _read_resume_text(row: UserPrefs) -> str:
    """Return raw text from uploaded resume for use as 'knowledge base'."""
    if row.resume_context:
        return row.resume_context
    if row.resume_file_path and Path(row.resume_file_path).exists():
        p = Path(row.resume_file_path)
        if p.suffix.lower() == ".pdf":
            return _extract_pdf_text(p)
        if p.suffix.lower() == ".docx":
            try:
                d = Document(str(p))
                return "\n".join(par.text for par in d.paragraphs if par.text.strip())
            except Exception:
                return ""
        return p.read_text(errors="ignore")
    return ""


def _extract_placeholders(docx_path: Path) -> list[str]:
    """Find {{name}} style placeholders inside the DOCX template paragraphs + tables."""
    import re
    from docx import Document as _D
    found = set()
    pattern = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")
    try:
        d = _D(str(docx_path))
        # paragraphs
        for p in d.paragraphs:
            for m in pattern.finditer(p.text or ""):
                found.add(m.group(1))
        # tables
        for tbl in d.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for m in pattern.finditer(p.text or ""):
                            found.add(m.group(1))
    except Exception as e:
        logger.warning("Placeholder extract failed: %s", e)
    return sorted(found)


from pydantic import BaseModel

class TailorJobRequest(BaseModel):
    job: dict     # {title, company, description, ...}


@router.post("/tailor-for-job")
def tailor_docx_for_job(req: TailorJobRequest, db: Session = Depends(get_db)):
    """Fill placeholders in user's uploaded DOCX template with JD-tailored content via Claude."""
    row = db.query(UserPrefs).first()
    if not row or not row.resume_file_path:
        raise HTTPException(status_code=400, detail="Upload a resume DOCX template first (Mission Config → Resume)")

    template_path = Path(row.resume_file_path)
    if not template_path.exists() or template_path.suffix.lower() != ".docx":
        raise HTTPException(status_code=400, detail="Resume must be a .docx template with {{placeholders}}")

    placeholders = _extract_placeholders(template_path)
    if not placeholders:
        raise HTTPException(
            status_code=400,
            detail="No {{placeholders}} found in your DOCX. Add placeholders like {{summary}}, {{role_1_bullets}} etc to your template.",
        )
    logger.info("Found %d placeholders in template: %s", len(placeholders), placeholders)

    knowledge = _read_resume_text(row)
    if not knowledge.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from resume DOCX")

    job = req.job
    jd = (job.get("description") or "")[:4000]
    role_target = job.get("title", "")
    company = job.get("company", "")

    placeholder_list = "\n".join(f"  - {p}" for p in placeholders)
    prompt = f"""You are tailoring a resume for a specific job. Fill in the placeholders below with JD-tailored content.

JOB:
Title: {role_target}
Company: {company}
Description:
{jd}

CANDIDATE RESUME (knowledge base — use these facts ONLY, never invent):
{knowledge[:6000]}

PLACEHOLDERS TO FILL (variable name → infer what content goes there from the name):
{placeholder_list}

RULES:
- Use only facts present in the resume above. Never invent companies, dates, metrics, or skills.
- Tailor wording to emphasize JD keywords where the underlying fact supports it.
- For bullet-list placeholders (e.g. role_1_bullets, project_X_bullets): return a single string with each bullet on its own line, NO bullet characters (the template handles bullet styling).
- For summary: 2-3 concise sentences highlighting fit for this specific job.
- For skills (e.g. skills_languages, skills_ai_ml): return comma-separated string of relevant skills only.
- For single-value placeholders (name, email, phone): return the value as a string.

Return ONLY valid JSON mapping each placeholder name to its filled value (string):
{{
{', '.join(f'"{p}": "..."' for p in placeholders[:5])}{', ...' if len(placeholders) > 5 else ''}
}}"""

    try:
        resp = get_client().chat.completions.create(
            model="anthropic/claude-sonnet-4.6",
            max_tokens=4000,
            messages=[{"role": "user", "content": prompt}],
        )
        text = resp.choices[0].message.content or ""
    except Exception as e:
        logger.error("Tailor LLM call failed: %s", e)
        raise HTTPException(status_code=502, detail=f"Tailor LLM failed: {e}")

    import re
    text = re.sub(r"```(?:json)?\s*", "", text).replace("```", "").strip()
    m = re.search(r"\{[\s\S]*\}", text)
    if not m:
        raise HTTPException(status_code=422, detail="Tailor LLM returned no JSON")
    try:
        context = json.loads(m.group())
    except json.JSONDecodeError as e:
        logger.error("JSON parse failed: %s\nLLM output: %s", e, m.group()[:500])
        raise HTTPException(status_code=422, detail=f"JSON parse failed: {e}")

    # Ensure every placeholder has SOME value (avoid Jinja UndefinedError)
    for ph in placeholders:
        if ph not in context:
            context[ph] = ""

    # Render template with docxtpl (preserves all formatting, fonts, layout)
    try:
        from docxtpl import DocxTemplate
        doc = DocxTemplate(str(template_path))
        doc.render(context)
        out_name = f"resume_{_slugify(company)}_{_slugify(role_target)}.docx"
        out_path = TAILORED_DIR / out_name
        doc.save(str(out_path))
    except Exception as e:
        logger.error("docxtpl render failed: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail=f"DOCX render failed: {e}")

    logger.info("Tailored DOCX written: %s (%d bytes)", out_path, out_path.stat().st_size)
    return {
        "file_path": str(out_path),
        "filename": out_name,
        "size": out_path.stat().st_size,
        "placeholders_filled": len(context),
        "company": company,
        "role": role_target,
    }


@router.get("/tailored/{filename}")
def download_tailored(filename: str):
    p = TAILORED_DIR / filename
    if not p.exists():
        raise HTTPException(status_code=404, detail="File not found")
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return FileResponse(str(p), media_type=media_type, filename=filename)


@router.get("/placeholders")
def list_placeholders(db: Session = Depends(get_db)):
    """Return the placeholder names found in user's uploaded DOCX template."""
    row = db.query(UserPrefs).first()
    if not row or not row.resume_file_path:
        return {"placeholders": [], "uploaded": False}
    p = Path(row.resume_file_path)
    if not p.exists() or p.suffix.lower() != ".docx":
        return {"placeholders": [], "uploaded": True, "is_docx": False}
    return {"placeholders": _extract_placeholders(p), "uploaded": True, "is_docx": True}


# ── Multi-resume library ─────────────────────────────────────────────────────
LIBRARY_DIR = UPLOAD_DIR / "library"
LIBRARY_DIR.mkdir(exist_ok=True)


def _resume_to_dict(r: Resume) -> dict:
    return {
        "id": r.id, "filename": r.filename, "role_tag": r.role_tag or "",
        "is_default": bool(r.is_default), "is_template": bool(r.is_template),
        "size": Path(r.file_path).stat().st_size if Path(r.file_path).exists() else 0,
        "created_at": r.created_at.isoformat() if r.created_at else None,
    }


@router.post("/library/upload")
async def library_upload(
    file: UploadFile = File(...),
    role_tag: str = Form(""),
    is_default: bool = Form(False),
    db: Session = Depends(get_db),
):
    """Upload a resume into the library, tagged by role (SDE / AI Engineer / FDE / ...)."""
    ext = Path(file.filename or "").suffix.lower()
    if ext not in ALLOWED_EXT:
        raise HTTPException(status_code=400, detail=f"Unsupported file type. Allowed: {ALLOWED_EXT}")

    # Unique filename in library dir (prefix with role_tag)
    role_slug = _slugify(role_tag) or "untagged"
    base = f"{role_slug}_{_slugify(Path(file.filename).stem)}{ext}"
    dest_path = LIBRARY_DIR / base
    n = 1
    while dest_path.exists():
        dest_path = LIBRARY_DIR / f"{role_slug}_{_slugify(Path(file.filename).stem)}_{n}{ext}"
        n += 1

    with dest_path.open("wb") as f:
        shutil.copyfileobj(file.file, f)

    text = ""
    if ext == ".pdf":
        text = _extract_pdf_text(dest_path)
    elif ext == ".docx":
        try:
            d = Document(str(dest_path))
            text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
        except Exception as e:
            logger.warning("docx extract failed: %s", e)
    elif ext == ".txt":
        text = dest_path.read_text(errors="ignore")

    is_template = ext == ".docx" and "{{" in text and "}}" in text

    if is_default:
        # Unset previous defaults
        for r in db.query(Resume).filter(Resume.is_default == 1).all():
            r.is_default = 0

    row = Resume(
        filename=file.filename or dest_path.name,
        file_path=str(dest_path.absolute()),
        role_tag=role_tag.strip(),
        extracted_text=text[:8000],
        is_default=1 if is_default else 0,
        is_template=1 if is_template else 0,
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    logger.info("Library resume uploaded: id=%d role=%r template=%s", row.id, row.role_tag, bool(row.is_template))
    return _resume_to_dict(row)


@router.get("/library")
def library_list(db: Session = Depends(get_db)):
    rows = db.query(Resume).order_by(Resume.is_default.desc(), Resume.created_at.desc()).all()
    return {"resumes": [_resume_to_dict(r) for r in rows]}


@router.delete("/library/{resume_id}")
def library_delete(resume_id: int, db: Session = Depends(get_db)):
    row = db.query(Resume).filter(Resume.id == resume_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="Resume not found")
    try: Path(row.file_path).unlink()
    except Exception: pass
    db.delete(row)
    db.commit()
    return {"ok": True}


@router.post("/library/{resume_id}/default")
def library_set_default(resume_id: int, db: Session = Depends(get_db)):
    row = db.query(Resume).filter(Resume.id == resume_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="Resume not found")
    for r in db.query(Resume).filter(Resume.is_default == 1).all():
        r.is_default = 0
    row.is_default = 1
    db.commit()
    return _resume_to_dict(row)


def _pick_best_resume(job_title: str, db: Session) -> Resume | None:
    """Pick the best matching resume from the library for a given job title.

    Strategy:
    1. Substring match on role_tag (case-insensitive, word-aware)
    2. Fall back to the default resume
    3. Fall back to the first non-template resume in library
    4. Fall back to None (caller uses generic upload path)
    """
    rows = db.query(Resume).all()
    if not rows:
        return None

    title_lower = (job_title or "").lower()
    title_words = set(re.findall(r"[a-z]+", title_lower))

    # Score each resume: count overlapping words between role_tag and job title
    best, best_score = None, 0
    for r in rows:
        if not r.role_tag:
            continue
        tag_words = set(re.findall(r"[a-z]+", r.role_tag.lower()))
        # Direct phrase match wins
        score = 100 if r.role_tag.lower() in title_lower else len(tag_words & title_words)
        if score > best_score:
            best, best_score = r, score

    if best:
        logger.info("Picked resume id=%d role_tag=%r for job=%r (score=%d)",
                    best.id, best.role_tag, job_title, best_score)
        return best

    default = next((r for r in rows if r.is_default), None)
    if default:
        logger.info("No role match — using default resume id=%d", default.id)
        return default

    logger.info("No role match + no default — using first library resume id=%d", rows[0].id)
    return rows[0]


@router.post("/library/pick-best")
def library_pick_best(payload: dict, db: Session = Depends(get_db)):
    """Pick best matching resume for a given job title. Body: {job_title: str}"""
    job_title = payload.get("job_title", "")
    r = _pick_best_resume(job_title, db)
    if not r:
        return {"resume": None}
    return {"resume": _resume_to_dict(r), "file_path": r.file_path}


@router.get("/library/{resume_id}/download")
def library_download(resume_id: int, db: Session = Depends(get_db)):
    row = db.query(Resume).filter(Resume.id == resume_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="Resume not found")
    p = Path(row.file_path)
    if not p.exists():
        raise HTTPException(status_code=404, detail="File missing on disk")
    return FileResponse(str(p), filename=row.filename)
